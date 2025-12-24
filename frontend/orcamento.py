import os
import sys
from datetime import date

import streamlit as st
from docx import Document

# Permite importar verificar_acesso
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from verificar_acesso import usuario_tem_acesso  # noqa: E402



def tela_orcamento(dados_empresa):
    """
    Tela principal de Orçamentos no frontend Streamlit.
    Gera DOCX estilizado + PDF (via reportlab) e lista orçamentos salvos.
    """
    # ⚠️ Verificação de acesso: Admin sempre tem acesso total
    email_usuario = st.session_state.get("dados_usuario", {}).get("email", "")
    if email_usuario != "matheus@email.com":
        if not usuario_tem_acesso("orcamento"):
            st.warning("⚠️ Este módulo está disponível apenas para planos pagos.")
            st.stop()

    st.title("🧾 Gerar Orçamento")

    # ------------------------------------------------------------------
    # DADOS DA EMPRESA
    # ------------------------------------------------------------------
    st.subheader("Dados da Empresa")
    col1, col2 = st.columns([1, 3])
    with col1:
        if dados_empresa.get("logo_url"):
            st.image(dados_empresa["logo_url"], width=150)
    with col2:
        st.markdown(f"**{dados_empresa.get('nome_empresa', '')}**")
        st.markdown(f"CNPJ: {dados_empresa.get('cnpj', '')}")
        st.markdown(
            f"{dados_empresa.get('rua', '')}, {dados_empresa.get('numero', '')} - "
            f"{dados_empresa.get('bairro', '')}"
        )
        st.markdown(
            f"{dados_empresa.get('cidade', '')} - CEP: {dados_empresa.get('cep', '')}"
        )

    st.markdown("---")

    # ------------------------------------------------------------------
    # DADOS DO CLIENTE
    # ------------------------------------------------------------------
    st.subheader("Dados do Cliente")
    nome_cliente = st.text_input("Nome do cliente")
    cpf_cnpj_cliente = st.text_input("CPF ou CNPJ do cliente")
    endereco_cliente = st.text_area("Endereço completo do cliente")

    st.markdown("---")

    # ------------------------------------------------------------------
    # ITENS DO ORÇAMENTO
    # ------------------------------------------------------------------
    st.subheader("Produtos ou Serviços")
    produtos = dados_empresa.get("produtos", [])
    lista_itens = []

    qtd_itens = st.number_input(
        "Quantos itens deseja incluir?", min_value=1, max_value=50, step=1, value=1
    )

    for i in range(qtd_itens):
        st.markdown(f"**Item {i + 1}**")
        col1, col2, col3 = st.columns(3)

        with col1:
            opcoes = [p.get("nome", "") for p in produtos] + ["Outro"]
            nome = st.selectbox("Produto/Serviço", opcoes, key=f"produto_{i}")
            if nome == "Outro":
                nome = st.text_input("Digite o nome do item", key=f"novo_produto_{i}")

        with col2:
            qtd = st.number_input(
                "Quantidade", min_value=1, step=1, key=f"qtd_{i}", value=1
            )

        with col3:
            valor = st.number_input(
                "Valor unitário (R$)",
                min_value=0.0,
                step=0.01,
                format="%.2f",
                key=f"valor_{i}",
            )

        lista_itens.append({"nome": nome, "qtd": qtd, "valor": valor})

    subtotal = sum(item["qtd"] * item["valor"] for item in lista_itens)
    desconto = st.number_input(
        "Desconto (R$)", min_value=0.0, step=0.01, format="%.2f"
    )
    total_final = max(subtotal - desconto, 0)

    st.markdown(
        f"**Subtotal: R$ {subtotal:,.2f}**"
        .replace(",", "X").replace(".", ",").replace("X", ".")
    )
    if desconto > 0:
        st.markdown(
            f"**Desconto:** R$ {desconto:,.2f}"
            .replace(",", "X").replace(".", ",").replace("X", ".")
        )
    st.markdown(
        f"### Total: R$ {total_final:,.2f}"
        .replace(",", "X").replace(".", ",").replace("X", ".")
    )

    st.markdown("---")

    # ------------------------------------------------------------------
    # INFORMAÇÕES COMPLEMENTARES
    # ------------------------------------------------------------------
    st.subheader("Informações Complementares")
    data_orcamento = st.date_input("Data do orçamento", value=date.today())
    validade = st.text_input("Prazo de validade do orçamento (ex: 10 dias)")
    prazo_execucao = st.text_input("Prazo para execução do serviço")
    observacoes = st.text_area("Observações adicionais")
    formas_pagamento = st.text_area(
        "Formas de pagamento (ex: parcelamento, desconto à vista)",
        value=(
            "- Parcelamento em até 12x no cartão, sem entrada;\n"
            "- Financiamento em até 48x via instituição financeira parceira;\n"
            "- Desconto especial para pagamento à vista."
        ),
    )

    if st.button("💾 Gerar Orçamento (DOCX e PDF)"):
        if not nome_cliente:
            st.error("Informe pelo menos o nome do cliente para gerar o orçamento.")
        else:
            salvar_orcamento_docx(
                dados_empresa,
                nome_cliente,
                cpf_cnpj_cliente,
                endereco_cliente,
                lista_itens,
                total_final,
                desconto,
                data_orcamento,
                validade,
                prazo_execucao,
                observacoes,
                formas_pagamento,
            )

    # ------------------------------------------------------------------
    # HISTÓRICO DE ORÇAMENTOS
    # ------------------------------------------------------------------
    st.markdown("---")
    st.subheader("📁 Orçamentos Salvos")

    dados_user = st.session_state.get("dados_usuario") or {}
    usuario_id = dados_user.get("id") or dados_user.get("usuario_id") or "anon"
    caminho = os.path.join("data", "clientes", str(usuario_id), "orcamentos")
    os.makedirs(caminho, exist_ok=True)

    arquivos = [f for f in os.listdir(caminho) if f.lower().endswith((".docx", ".pdf"))]
    if not arquivos:
        st.info("Nenhum orçamento salvo até o momento.")
        return

    # Filtros
    colf1, colf2 = st.columns(2)
    with colf1:
        nome_filtro = st.text_input("🔍 Filtrar por nome do cliente")
    with colf2:
        data_filtro = st.date_input(
            "📅 Filtrar por data do orçamento", value=None, key="filtro_data_orc"
        )

    registros = []
    for nome_arquivo in arquivos:
        base, ext = os.path.splitext(nome_arquivo)
        partes = base.split("_")

        cliente = "Desconhecido"
        data_str = ""
        if len(partes) >= 3:
            # orcamento_nome_cliente_YYYYMMDD
            data_str = partes[-1]
            cliente = " ".join(partes[1:-1]).replace("-", " ")
        elif len(partes) == 2:
            cliente = partes[1]

        data_fmt = None
        if len(data_str) == 8 and data_str.isdigit():
            ano = int(data_str[:4])
            mes = int(data_str[4:6])
            dia = int(data_str[6:])
            try:
                data_fmt = date(ano, mes, dia)
            except Exception:
                data_fmt = None

        registros.append(
            {
                "arquivo": nome_arquivo,
                "cliente": cliente,
                "data": data_fmt,
                "ext": ext.lower().strip("."),
            }
        )

    # Aplica filtros
    registros_filtrados = []
    for r in registros:
        if nome_filtro and nome_filtro.lower() not in r["cliente"].lower():
            continue
        if data_filtro and r["data"] and r["data"] != data_filtro:
            continue
        registros_filtrados.append(r)

    if not registros_filtrados:
        st.info("Nenhum orçamento encontrado com os filtros atuais.")
        return

    # Ordena: mais recentes primeiro
    registros_ordenados = sorted(
        registros_filtrados,
        key=lambda r: (r["data"] or date.min, r["arquivo"]),
        reverse=True,
    )

    for reg in registros_ordenados:
        arq = reg["arquivo"]
        cliente = reg["cliente"]
        data_fmt = reg["data"]
        data_str = data_fmt.strftime("%d/%m/%Y") if data_fmt else "Data não identificada"

        col1, col2, col3 = st.columns([3, 2, 2])
        with col1:
            st.markdown(f"**{cliente}**  \n📄 `{arq}`  \n🗓 {data_str}")

        full_path = os.path.join(caminho, arq)
        with col2:
            if os.path.exists(full_path):
                with open(full_path, "rb") as f:
                    st.download_button(
                        f"⬇️ Baixar {reg['ext'].upper()}",
                        data=f,
                        file_name=arq,
                        mime="application/octet-stream",
                        key=f"down_{arq}",
                    )
        with col3:
            link = f"https://wa.me/?text=Segue%20seu%20orçamento:%20{arq}"
            st.markdown(f"[📲 Enviar por WhatsApp]({link})")


# ======================================================================
# FUNÇÃO QUE GERA DOCX + PDF (via REPORTLAB)
# ======================================================================
def salvar_orcamento_docx(
    empresa,
    nome_cliente,
    cpf_cnpj,
    endereco_cliente,
    itens,
    total,
    desconto,
    data_orcamento,
    validade,
    prazo_execucao,
    observacoes,
    formas_pagamento,
):
    """
    Gera um DOCX profissional (cabeçalho, tabela, blocos) e,
    em seguida, gera um PDF separado usando reportlab.
    NÃO usa docx2pdf.
    """

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Inches, Pt

    def format_brl(valor: float) -> str:
        return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    # ------------------------------------------------------------------
    # GERAR DOCX
    # ------------------------------------------------------------------
    doc = Document()

    # Margens da página
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # CABEÇALHO
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = True
    cell_left, cell_right = header_table.rows[0].cells

    # Empresa à esquerda
    p_emp = cell_left.paragraphs[0]
    p_emp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_emp = p_emp.add_run(empresa.get("nome_empresa", ""))
    run_emp.bold = True
    run_emp.font.size = Pt(16)

    p_emp2 = cell_left.add_paragraph()
    p_emp2.add_run(f"CNPJ: {empresa.get('cnpj', '')}\n").font.size = Pt(9)
    p_emp2.add_run(
        f"{empresa.get('rua', '')}, {empresa.get('numero', '')} - {empresa.get('bairro', '')}\n"
        f"{empresa.get('cidade', '')} - CEP {empresa.get('cep', '')}"
    ).font.size = Pt(9)

    # Título ORÇAMENTO à direita
    p_orc = cell_right.paragraphs[0]
    p_orc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_title = p_orc.add_run("ORÇAMENTO")
    run_title.bold = True
    run_title.font.size = Pt(18)

    p_orc2 = cell_right.add_paragraph()
    p_orc2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_data = p_orc2.add_run(data_orcamento.strftime("%d/%m/%Y"))
    run_data.font.size = Pt(9)

    doc.add_paragraph("")

    # DADOS DO CLIENTE
    p_cli_title = doc.add_paragraph()
    p_cli_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_cli_title = p_cli_title.add_run("Dados do Cliente")
    run_cli_title.bold = True
    run_cli_title.font.size = Pt(11)

    p_cli = doc.add_paragraph()
    p_cli.add_run("Nome: ").bold = True
    p_cli.add_run(nome_cliente or "")

    p_cpf = doc.add_paragraph()
    p_cpf.add_run("CPF/CNPJ: ").bold = True
    p_cpf.add_run(cpf_cnpj or "")

    p_end = doc.add_paragraph()
    p_end.add_run("Endereço: ").bold = True
    p_end.add_run(endereco_cliente or "")

    doc.add_paragraph("")

    p_ap = doc.add_paragraph()
    p_ap.add_run(
        "Apresentamos nossa proposta exclusiva com os itens abaixo descritos:"
    ).font.size = Pt(10)

    doc.add_paragraph("")

    # TABELA DE ITENS
    tabela = doc.add_table(rows=1, cols=4)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.style = "Light List Accent 1"

    hdr_cells = tabela.rows[0].cells
    cabecalhos = ["Descrição", "Qtd", "Valor unitário", "Valor total"]
    for idx, texto in enumerate(cabecalhos):
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(texto)
        run.bold = True
        run.font.size = Pt(10)

    subtotal = 0.0
    for item in itens:
        nome = item.get("nome", "")
        qtd = int(item.get("qtd") or 0)
        valor = float(item.get("valor") or 0.0)
        total_item = qtd * valor
        subtotal += total_item

        row_cells = tabela.add_row().cells
        row_cells[0].text = str(nome)
        row_cells[1].text = str(qtd)
        row_cells[2].text = format_brl(valor)
        row_cells[3].text = format_brl(total_item)

        for idx in range(4):
            for p in row_cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph("")

    # RESUMO
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sub.add_run(f"Subtotal: {format_brl(subtotal)}").font.size = Pt(10)

    if desconto and desconto > 0:
        p_desc = doc.add_paragraph()
        p_desc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_desc.add_run(f"Desconto: {format_brl(desconto)}").font.size = Pt(10)

    p_tot = doc.add_paragraph()
    p_tot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_tot = p_tot.add_run(f"Total: {format_brl(total)}")
    run_tot.bold = True
    run_tot.font.size = Pt(14)

    doc.add_paragraph("")

    # BLOCO INFORMAÇÕES
    info_table = doc.add_table(rows=2, cols=3)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = True

    headers_info = ["Validade da Proposta", "Prazo de Execução", "Data do Orçamento"]
    for i, texto in enumerate(headers_info):
        p = info_table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(texto)
        r.bold = True
        r.font.size = Pt(9)

    val_row = info_table.rows[1].cells
    val_row[0].text = validade or ""
    val_row[1].text = prazo_execucao or ""
    val_row[2].text = data_orcamento.strftime("%d/%m/%Y")

    for cell in val_row:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(9)

    doc.add_paragraph("")

    # OBSERVAÇÕES
    p_obs_title = doc.add_paragraph()
    run_obs_title = p_obs_title.add_run("Observações da Proposta")
    run_obs_title.bold = True
    run_obs_title.font.size = Pt(11)

    if observacoes:
        p_obs = doc.add_paragraph(observacoes)
        for r in p_obs.runs:
            r.font.size = Pt(9)
    else:
        doc.add_paragraph("")

    doc.add_paragraph("")

    # FORMAS DE PAGAMENTO
    p_fp_title = doc.add_paragraph()
    run_fp_title = p_fp_title.add_run("Formas de Pagamento")
    run_fp_title.bold = True
    run_fp_title.font.size = Pt(11)

    if formas_pagamento:
        for linha in formas_pagamento.splitlines():
            if linha.strip():
                p_fp = doc.add_paragraph(linha.strip())
                p_fp.style = doc.styles["List Bullet"]
                for r in p_fp.runs:
                    r.font.size = Pt(9)
    else:
        doc.add_paragraph("")

    doc.add_paragraph("")

    # ASSINATURAS
    doc.add_paragraph("\n")
    assinatura_table = doc.add_table(rows=1, cols=3)
    assinatura_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    assinatura_table.autofit = True

    campos = ["Nome Legível", "Assinatura", "Data"]
    for i, texto in enumerate(campos):
        cell = assinatura_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("__________________________")
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(texto).font.size = Pt(8)

    # --------------------------------------------------------------
    # SALVAR DOCX
    # --------------------------------------------------------------
    slug_cliente = nome_cliente.replace(" ", "_").lower() if nome_cliente else "cliente"
    nome_base = f"orcamento_{slug_cliente}_{data_orcamento.strftime('%Y%m%d')}"
    dados_user = st.session_state.get("dados_usuario") or {}
    usuario_id = dados_user.get("id") or dados_user.get("usuario_id") or "anon"

    caminho = os.path.join("data", "clientes", str(usuario_id), "orcamentos")
    os.makedirs(caminho, exist_ok=True)

    docx_path = os.path.join(caminho, f"{nome_base}.docx")
    pdf_path = os.path.join(caminho, f"{nome_base}.pdf")

    doc.save(docx_path)

    # --------------------------------------------------------------
    # GERAR PDF VIA REPORTLAB
    # --------------------------------------------------------------
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib import colors
        from reportlab.platypus import Table, TableStyle

        width, height = A4
        c = canvas.Canvas(pdf_path, pagesize=A4)

        y = height - 60

        # Cabeçalho
        c.setFont("Helvetica-Bold", 14)
        c.drawString(40, y, empresa.get("nome_empresa", ""))
        y -= 16
        c.setFont("Helvetica", 9)
        c.drawString(40, y, f"CNPJ: {empresa.get('cnpj', '')}")
        y -= 12
        endereco_emp = (
            f"{empresa.get('rua', '')}, {empresa.get('numero', '')} - "
            f"{empresa.get('bairro', '')}, {empresa.get('cidade', '')} - CEP {empresa.get('cep', '')}"
        )
        c.drawString(40, y, endereco_emp[:110])
        y -= 22

        c.setFont("Helvetica-Bold", 18)
        c.drawRightString(width - 40, height - 60, "ORÇAMENTO")
        c.setFont("Helvetica", 9)
        c.drawRightString(width - 40, height - 80, data_orcamento.strftime("%d/%m/%Y"))

        # Dados do cliente
        y -= 10
        c.setFont("Helvetica-Bold", 11)
        c.drawString(40, y, "Dados do Cliente")
        y -= 14
        c.setFont("Helvetica", 9)
        c.drawString(40, y, f"Nome: {nome_cliente}")
        y -= 12
        c.drawString(40, y, f"CPF/CNPJ: {cpf_cnpj}")
        y -= 12
        c.drawString(40, y, f"Endereço: {endereco_cliente[:110]}")
        y -= 20

        c.setFont("Helvetica", 9)
        c.drawString(
            40,
            y,
            "Apresentamos nossa proposta exclusiva com os itens abaixo descritos:",
        )
        y -= 18

        # Tabela itens
        data_table = [["Descrição", "Qtd", "Unitário", "Total"]]
        for item in itens:
            nome = str(item.get("nome", ""))
            qtd = int(item.get("qtd") or 0)
            valor = float(item.get("valor") or 0.0)
            total_item = qtd * valor
            data_table.append(
                [nome, str(qtd), format_brl(valor), format_brl(total_item)]
            )

        table = Table(data_table, colWidths=[220, 40, 80, 80])
        style = TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 9),
                ("FONTSIZE", (0, 1), (-1, -1), 8),
                ("ALIGN", (1, 1), (-1, -1), "CENTER"),
            ]
        )
        table.setStyle(style)

        w, h = table.wrapOn(c, width - 80, y)
        table.drawOn(c, 40, y - h)
        y = y - h - 20

        # Resumo valores
        c.setFont("Helvetica", 9)
        c.drawRightString(width - 40, y, f"Subtotal: {format_brl(subtotal)}")
        y -= 12
        if desconto and desconto > 0:
            c.drawRightString(width - 40, y, f"Desconto: {format_brl(desconto)}")
            y -= 12
        c.setFont("Helvetica-Bold", 12)
        c.drawRightString(width - 40, y, f"Total: {format_brl(total)}")
        y -= 24

        # Informações da proposta
        c.setFont("Helvetica-Bold", 10)
        c.drawString(40, y, "Informações da Proposta")
        y -= 14
        c.setFont("Helvetica", 9)
        c.drawString(40, y, f"Validade da proposta: {validade}")
        y -= 12
        c.drawString(40, y, f"Prazo de execução: {prazo_execucao}")
        y -= 12
        c.drawString(40, y, f"Data do orçamento: {data_orcamento.strftime('%d/%m/%Y')}")
        y -= 18

        # Observações
        if observacoes:
            c.setFont("Helvetica-Bold", 10)
            c.drawString(40, y, "Observações da Proposta")
            y -= 14
            c.setFont("Helvetica", 9)
            for linha in observacoes.splitlines():
                c.drawString(40, y, linha[:110])
                y -= 12
                if y < 80:
                    c.showPage()
                    y = height - 60

        # Formas de pagamento
        if y < 120:
            c.showPage()
            y = height - 60

        c.setFont("Helvetica-Bold", 10)
        c.drawString(40, y, "Formas de Pagamento")
        y -= 14
        c.setFont("Helvetica", 9)
        for linha in (formas_pagamento or "").splitlines():
            if linha.strip():
                c.drawString(40, y, linha[:110])
                y -= 12
                if y < 80:
                    c.showPage()
                    y = height - 60

        # Assinaturas
        if y < 100:
            c.showPage()
            y = height - 60

        y_assin = 80
        c.setFont("Helvetica", 9)
        c.line(60, y_assin + 10, 180, y_assin + 10)
        c.drawCentredString(120, y_assin, "Nome Legível")

        c.line(240, y_assin + 10, 360, y_assin + 10)
        c.drawCentredString(300, y_assin, "Assinatura")

        c.line(420, y_assin + 10, 540, y_assin + 10)
        c.drawCentredString(480, y_assin, "Data")

        c.save()

    except ImportError:
        st.warning(
            "PDF não gerado: biblioteca 'reportlab' não está instalada. "
            "Instale com: pip install reportlab"
        )
    except Exception as e:
        st.warning(f"Erro ao gerar PDF: {e}")

    # ------------------------------------------------------------------
    # BOTÕES DE DOWNLOAD
    # ------------------------------------------------------------------
    st.success("Orçamento salvo com sucesso.")
    with open(docx_path, "rb") as f_docx:
        st.download_button(
            "📥 Baixar DOCX",
            data=f_docx,
            file_name=os.path.basename(docx_path),
        )
    if os.path.exists(pdf_path):
        with open(pdf_path, "rb") as f_pdf:
            st.download_button(
                "📥 Baixar PDF",
                data=f_pdf,
                file_name=os.path.basename(pdf_path),
            )
